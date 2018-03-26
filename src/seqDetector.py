#!/usr/bin/python
# coding: utf-8
# Write by Richard Bonnet
# Date: 19/01/2016
import fnmatch
import os
import argparse
from Bio import SeqIO
import datetime


install_dir = os.path.dirname(os.path.realpath(__file__))

def recursive_fileList(directory, filter):
    fileList = []
    for root, dirnames, filenames in os.walk(directory):
        for filename in fnmatch.filter(filenames, filter):
            fileList.append(os.path.join(root, filename))
    return fileList


def read_sample_file(sampleFile):
    sampleDict = {}
    f = open(sampleFile,'r')
    for line in f:
        line = line.strip()
        if line != '' and line.startswith('#') == False:
            sample_id = line.split('\t')[0]
            species = line.split('\t')[1]  
            sampleDict[sample_id] = species.lower()
    return sampleDict


def rename_ctg(fastaFile):
    recList = []
    f = open(fastaFile, 'r')
    for i, rec in enumerate(SeqIO.parse(f,'fasta')):
        rec.id = 'ctg_%i' % (i+1)
        rec.description = ''
        recList.append(rec)
    f.close()
    if os.path.exists(os.path.splitext(fastaFile)[0] + '_fasta') == False:
        cmd = 'cp %s %s' % (fastaFile, os.path.splitext(fastaFile)[0] + '_contigs.assembly')
        os.system(cmd)
    f = open(fastaFile, 'w')
    SeqIO.write(recList, f, 'fasta')
    f.close()


def runProkka(sample_id, fastaFile, genus, species, evalue):
    cmd = '%s -f %s -g %s -s %s -e %s -id %s' % (os.path.join(install_dir, 'runProkka.py'), fastaFile, genus, species, evalue, sample_id)
    #print cmd
    os.system(cmd)


def runMLST(fasFile, outdir, sampleID, sampleFile):
    cmd = '%s -f %s -o %s -id %s -s %s ' % (os.path.join(install_dir, 'seqMLST.py'), fasFile, outdir, sampleID, sampleFile)
    os.system(cmd)


def setupblastDB(csvdbName, catDB, dbType, force):
    outdbFasta = os.path.splitext(csvdbName)[0]
    if dbType == 'dna':
        outdbFasta = outdbFasta + '.fna'
    else:
        outdbFasta = outdbFasta + '.faa'
    if os.path.exists(outdbFasta) == False or force == True:
        cmd = '%s -cnr %s -cat %s -o %s -t %s ' % (os.path.join(install_dir, 'cnrdbTofasta.py'), csvdbName, catDB, os.path.splitext(csvdbName)[0], dbType)
        #print cmd
        os.system(cmd)
    else:
        print '\nFasta DB file %s already exists!\n' % outdbFasta


def runBlast(fasFile, fasdbName, outdir, blastType, threads, blastevalue, force):

    xmlFile = os.path.join(outdir, os.path.basename(fasFile) + '_' + os.path.splitext(os.path.basename(fasdbName))[0] + '.blastRes.xml')
    csvFile = xmlFile + '.csv'
    if os.path.exists(csvFile) == False or force == True:
        cmd = '%s -q %s -db %s -o %s -bt %s -t %s -e %s' % (os.path.join(install_dir, 'runBlast.py'), fasFile, fasdbName, outdir, blastType, threads, blastevalue)
        #print cmd
        os.system(cmd)
    else:
        print '\nBlast parsing file %s detected\nBlast already done!\n' % csvFile
        xmlFile = ''
    return xmlFile, csvFile


def parseBlast(fasFile, csvdbName, blastresFile, outdir, sequenceType):
    cmd = '%s -q %s -db %s -xml %s -o %s -t %s ' % (os.path.join(install_dir, 'parseBlast.py'), fasFile, csvdbName, blastresFile, outdir, sequenceType)
    #print cmd
    os.system(cmd)
    cmd = 'rm %s' % blastresFile
    os.system(cmd)


def mergeResults(workdir, idperc, cov, outdir):
    cmd = '%s -d %s -id %.2f -cv %.2f -o %s' % \
          (os.path.join(install_dir, 'mergeResults.py'), workdir, float(idperc), float(cov), outdir)
    #print cmd
    os.system(cmd)


def annotGBK(gbkFile, csvFile, idperc, cov, outdir):
    cmd = '%s -gbk %s -csv %s -id %.2f -cv %.2f -o %s' % \
          (os.path.join(install_dir, 'annotGBK.py'), gbkFile, csvFile, float(idperc), float(cov), outdir)
    print '\nEdition of gbk file %s' % gbkFile
    os.system(cmd)
    print 'Edition of gbk file %s done!\n' % gbkFile


def write_docx(csvblastresFile, idperc, cov, wkdir, sample_id, sampleFile, csvdbName, catDB, initial):
    f = open(sampleFile, 'r')
    for line in f:
        if line.split('\t')[0] == sample_id:
	    #Aurelien : j ai rajouter le lowercase
            species = line.strip().split('\t')[1].lower()
            break
    f.close()

    found = 1
    mlstFile = os.path.join(install_dir, 'db_mlst', 'mlst.txt')
    f = open(mlstFile, 'r')
    for line in f:
        if line.split('\t')[0].lower() == species.lower():
            mlst_name = line.strip().split('\t')[1]
            found = 0
            break
    f.close()

    if found == 0:

        mlstFile = os.path.join(wkdir, species.replace(' ', '_') + '__' + mlst_name + '__MLST.csv')
        f = open(mlstFile, 'r')
        ST = ''
        for line in f:
            line = line.strip()
            if line.split('\t')[0] == sample_id:
                ST = line.split('\t')[1]
                break
        f.close()
    else:
        ST = ''

    if ST == '?':
        ST = ''

    print '\nData from %s' % csvblastresFile
    key = sample_id
    fresDict={}
    f = open(csvblastresFile, 'r')
    for n, line in enumerate(f):
        line = line.strip()
        if n == 0:
            header = line.split('\t')
        else:
            data = line.split('\t')
            dic = dict(zip(header, data))

            fCluster = dic['Functional cluster'].replace('|', '-')

            if fCluster.strip() == '':
                fCluster = 'Undefined-Cluster'

            if float(dic['Identity (%)']) >= float(idperc) and \
                            float(dic['Hit coverage (%)']) >= float(cov) and \
                            float(dic['Query coverage (%)']) >= float(cov):
                res = ''
                if dic['SNP(s) searched'] == '':
                    if float(dic['Identity (%)']) == 100 and float(dic['Hit coverage (%)']) == 100 and float(dic['Query coverage (%)']) == 100:
                        res = dic['Hit name']
                    elif float(dic['Hit coverage (%)']) == 100 and dic['Hit dtbase protein sequence'] in dic['Query sequence']:
                        res = dic['Hit name']
                    elif float(dic['Query coverage (%)']) == 100 and dic['Query sequence'] in dic['Hit dtbase protein sequence']:
                        res = dic['Hit name']
                    else:
                        res = dic['Hit name'] + '-like'
                elif dic['SNP(s) searched'] != '' and dic['SNP(s) found'] != '':
                    res = 'variant resistant de '.decode('utf-8') + '%s' % dic['Hit name'].strip().split(' ')[0]

                if species in ['klebsiella pneumoniae', 'enterobacter cloacae'] and 'oqx' in res:
                    res = ''

                if res == 'AmpC [E. coli]-like':
                    res = 'AmpC [E. coli]'

                #FUNCTIONAL CLASSIFICATION
                try:
                    if res != '':
                        fresDict[fCluster].append(res)
                except KeyError:
                    if res != '':
                        fresDict[fCluster] = [res]
                        fresDict[fCluster] = list(set(fresDict[fCluster]))

    from docx import Document
    from docx.shared import Inches

    document = Document()
    sections = document.sections
    sections[0].left_margin = Inches(0.5)
    sections[0].right_margin = Inches(0.5)

    ref = document.add_heading('Référence CNR de l\'isolat bacterien : \t'.decode('utf-8'), 3)
    ref.add_run(sample_id).italic = True

    spp = document.add_heading('Espèce bactérienne :\t'.decode('utf-8'), 3)
    spp.add_run(species[0].upper() + species[1:] + '\n').italic = True

    document.add_heading('Matériel & Méthodes :'.decode('utf-8'), 3)
    met_p1 = document.add_paragraph('   ~ Séquençage : '.decode('utf-8'))
    met_p1.add_run('Méthode Illumina (2 x 150pb ou 300pb appariés)\n\n'.decode('utf-8'))
    met_p1.add_run('   ~ Analyse '.decode('utf-8'))
    met_p1.add_run('in silico '.decode('utf-8')).italic = True
    met_p1.add_run('des séquences génomiques : prokka, blastn, blastp et pairwise2\n\n'.decode('utf-8'))
    if ST != '':
        met_p1.add_run('   ~ Bases de données MLST : '.decode('utf-8'))
        if species == 'escherichia coli':
            met_p1.add_run('http://mlst.warwick.ac.uk/mlst\n\n'.decode('utf-8'))
        elif species == 'klebsiella pneumoniae':
            met_p1.add_run('http://bigsdb.pasteur.fr\n\n'.decode('utf-8'))
        else:
            met_p1.add_run('https://pubmlst.org/databases\n\n'.decode('utf-8'))

    met_p1.add_run('   ~ Bases de données CNR de la résistance aux antibiotiques : %s cat.: %s\n'.decode('utf-8')
                   % (os.path.splitext(os.path.basename(csvdbName))[0].replace('_', ' v.: '), catDB))

    if ST != '':
        document.add_heading('Résultat : Génotypage MLST '.decode('utf-8'), 3)
        document.add_paragraph('   ~ Sequence Type: ST-%s\n'.decode('utf-8') % ST)

    document.add_heading('Résultat : Déterminants de la résistance aux 3 principales familles d\'antibiotiques (*)'.decode('utf-8'), 3)


    frDict = {'Aminoglycoside':"Aminosides",'Beta-lactam':"Beta-lactamines".decode('utf-8'),'Quinolone':"Quinolones",
              'Sulfonamide':"Sulfamides",'Trimethoprime':"Triméthoprime".decode('utf-8'),'Cycline':"Tétracycline".decode('utf-8')}
    fresults = ['Aminoglycoside','Beta-lactam','Quinolone'] #,'Sulfonamide','Trimethoprime','Cycline']
    for func in fresults:
        res_txt = ''
        try:
            res = fresDict[func]
            res.sort()
            res_txt = res_txt + '   ~ %s : %s\n'.decode('utf-8') % (frDict[func], ', '.join(res))
        except KeyError:
            res_txt = res_txt + '   ~ %s :\n'.decode('utf-8') % (frDict[func])
        document.add_paragraph(res_txt.decode('utf-8'))
    print res_txt
    document.add_paragraph('(*) Contacter le CNR pour d\'autres familles d\'antibiotiques (tél: 04 73 754 920)'.decode('utf-8'))
    document.save(os.path.join(wkdir, 'CR_CNR_%s_%s_%s.docx' % (sample_id, datetime.date.today(), initial)))


def main(args):

    force = args.force
    initial = args.initial
    csvdbName = args.cnrDB
    catDB = args.catDB
    sequenceType = args.seqType
    if sequenceType == '1' or sequenceType == '2':
        dbType = 'dna'
        fasdbPrefix = os.path.splitext(csvdbName)[0] + '_' + catDB
        fasdbName = fasdbPrefix + '.fna'
        blastType = 'blastn'

    else:
        dbType = 'prot'
        fasdbPrefix = os.path.splitext(csvdbName)[0]  + '_' + catDB
        fasdbName = fasdbPrefix + '.faa'
        blastType = 'blastp'

    setupblastDB(csvdbName, catDB, dbType, force)

    threads = args.threads
    blastevalue = args.evalue
    idperc = args.idPerc
    cov = args.Cov

    wkdir = args.wkDir
    contigExt = args.contigExt
    sampleFile = args.sampleFile
    if sampleFile == '':
        sampleFile = os.path.join(wkdir, 'sample.csv')
    if os.path.exists(sampleFile) == True:
        print '\nThe sample file %s have been found.\n' % sampleFile
        sampleDict = read_sample_file(sampleFile)
        sampleList = sampleDict.keys()
        sampleList.sort()
        for sample_id in sampleList:
            print '\n\nStart %s' % sample_id
            gbkFiles = recursive_fileList(wkdir, sample_id + '.gbk')
            gbkFiles.sort()
            faaFiles = recursive_fileList(wkdir, sample_id + '.faa')
            faaFiles.sort()
            
            import datetime
            now = str(datetime.datetime.now())
            f = open('/usr/local/seqblast/cnr/sdet.txt','a')
            f.write(sample_id + '\t' + now + '\t' + initial +'\n')
            f.close()

            if gbkFiles != [] and faaFiles != [] \
            and os.path.dirname(gbkFiles[0]) == os.path.dirname(faaFiles[0]) and force == False:
                outdir = os.path.dirname(gbkFiles[0])
                print '\nAnnotation already done in %s!\n' % outdir
            else:
                fastaFiles = recursive_fileList(wkdir, sample_id + '*.' + contigExt)
                fastaFiles.sort()
                fastaFile = fastaFiles[0]
                rename_ctg(fastaFile)
                genus, species = sampleDict[sample_id].split(' ')
                runProkka(sample_id, fastaFile, genus, species, '1e-30')
                outdir = os.path.dirname(fastaFile)

            fastaFiles = recursive_fileList(wkdir, sample_id + '*.fasta')
            fastaFiles.sort()
            runMLST(fastaFiles[0], wkdir, sample_id, sampleFile)

            if sequenceType == '0':
                fasFile = os.path.join(outdir, sample_id + '.faa')
            elif sequenceType == '1':
                fasFile = os.path.join(outdir, sample_id + '.fna')
            elif sequenceType == '2':
                fasFile = os.path.join(outdir, sample_id + '.fsa')
            gbkFile = os.path.join(outdir, sample_id + '.gbk')

            xmlblastresFile, csvblastresFile = runBlast(fasFile, fasdbName, outdir, blastType, threads, blastevalue, force)
            if xmlblastresFile != '':
                parseBlast(fasFile, csvdbName, xmlblastresFile, outdir, sequenceType)
            write_docx(csvblastresFile, idperc, cov, wkdir, sample_id, sampleFile, csvdbName, catDB, initial)
            annotGBK(gbkFile, csvblastresFile, idperc, cov, outdir)
        mergeResults(wkdir, idperc, cov, wkdir)
    else:
        print '\nNo sample file!\n'
        print exit()

def version():
    return "1.0"


def run():
    parser = argparse.ArgumentParser(description='seqDetector - Version ' + version())
    parser.add_argument('-d','--wkDir', dest="wkDir", default='/home/bacteriologie/assemblage/wk',help='The work directory containing the data')
    parser.add_argument('-ext','--contigExt', dest="contigExt", default="fasta", help="The extension of contig files (default: fasta)")
    parser.add_argument('-s','--sampleFile', dest="sampleFile", default='', help="The tab sample file with corresponding bacterial species (default: $wkDir/sample.csv)")
    parser.add_argument('-db','--cnrDB', dest="cnrDB", default='/usr/local/seqblast/cnr/armDB_1.csv', help="A CNR database as csv file (default: usr/local/seqblast/cnr/armDB_1.csv)")
    parser.add_argument('-cat','--catDB', dest="catDB", default='GN', help="Category of record in CNR database (default: GN)")
    parser.add_argument('-st','--seqType', dest="seqType", default='0', help="0 for protein (blastp), 1 for CDS (blastn), 2 for contigs (blastn) (default: 0)")
    parser.add_argument('-e','--evalue', dest="evalue", default='1e-30', help="blast evalue cutoff (default: 1e-30)")
    parser.add_argument('-t','--threads', dest="threads", default='8', help="Thread number for blast (default: 8)")
    parser.add_argument('-id','--idPerc', dest="idPerc", default='80', help="Identity percentage cutoff for annotation (default: 80)")
    parser.add_argument('-cv','--cov', dest="Cov", default='80', help="Coverage percentage cutoff for annotation (default: 80)")
    parser.add_argument('-in','--initial', dest="initial", help="Initials of the operator")
    parser.add_argument('-F','--force', dest="force", action='store_true',default=False, help="Force file overwrite (default: False)")
    parser.add_argument('-V', '--version', action='version', version='seqDetector-' + version(), help="Prints version number")
    args = parser.parse_args()
    main(args)


if __name__ == '__main__':
    run()
