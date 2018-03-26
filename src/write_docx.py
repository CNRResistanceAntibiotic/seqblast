#!/usr/bin/python
# coding: utf-8
from docx import Document
from docx.shared import Inches

document = Document()
sections = document.sections
sections[0].left_margin = Inches(1)
sections[0].right_margin =Inches(1)

sample_id = 'CNRxxx'
ref = document.add_heading('Référence CNR de l\'isolat bacterien : \t'.decode('utf-8'), 2)
ref.add_run(sample_id).italic = True

species = 'Escherichia coli'
spp = document.add_heading('Espèce bactérienne : \t'.decode('utf-8'), 2)
spp.add_run(species).italic = True

met = document.add_heading('Matériel & Méthodes :'.decode('utf-8'), 2)
met_p1 = document.add_paragraph('   * Séquençage : '.decode('utf-8'))
met_p1.add_run('Méthode Illumina (2 x 150pb ou 300pb appariés)\n'.decode('utf-8'))
met_p1.add_run('   * Analyse '.decode('utf-8'))
met_p1.add_run('in silico '.decode('utf-8')).italic = True
met_p1.add_run('des séquences génomiques : prokka, blastn, blastp et pairwise2\n'.decode('utf-8'))
met_p1.add_run('   * Bases de données MLST :\n'.decode('utf-8'))
met_p1.add_run('\t- Escherichia coli'.decode('utf-8')).italic = True
met_p1.add_run(' : http://mlst.warwick.ac.uk/mlst\n'.decode('utf-8'))
met_p1.add_run('\t- Klebsiella spp.'.decode('utf-8')).italic = True
met_p1.add_run(' : http://bigsdb.pasteur.fr\n'.decode('utf-8'))
met_p1.add_run('\t- Autres espèces : https://pubmlst.org/databases\n'.decode('utf-8'))
met_p1.add_run('   * Bases de données des déterminants de la résistance aux antibiotiques :\n'.decode('utf-8'))
met_p1.add_run('\t- Base de données CNR (2700 déterminants)'.decode('utf-8'))

res_mlst = document.add_heading('Résultat : Génotypage MLST '.decode('utf-8'), 2)

res_amr = document.add_heading('Résultat : Déterminants de la résistance aux antibiotiques'.decode('utf-8'), 2)

document.save('essai.docx')